#!/usr/bin/env python3
"""
ApexEval - AI-Powered Candidate Evaluation Platform
Built for Apexire | Intelligent Hiring
Author: Grok for Saulo Borges / Apexire
"""

import streamlit as st
import os
import io
import json
import zipfile
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
import base64

from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from PIL import Image
import pdfplumber
from pypdf import PdfReader

# ====================== CONFIG ======================
APP_NAME = "ApexEval"
APP_TAGLINE = "AI-Powered Candidate Evaluation | Powered by Apexire"
APEXIRE_BLUE = RGBColor(0, 150, 220)  # Primary brand color
APEXIRE_DARK = RGBColor(20, 40, 80)
SUCCESS_GREEN = RGBColor(40, 167, 69)
WARNING_YELLOW = RGBColor(255, 193, 7)
DANGER_RED = RGBColor(220, 53, 69)

DEFAULT_MODEL = "grok-3"
MAX_CVS = 10

# ====================== PAGE SETUP ======================
st.set_page_config(
    page_title=f"{APP_NAME} | Apexire",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional Apexire look
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #003366 0%, #0096DC 100%);
        padding: 1.5rem 2rem;
        border-radius: 12px;
        color: white;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 12px rgba(0, 150, 220, 0.3);
    }
    .section-header {
        color: #003366;
        border-bottom: 3px solid #0096DC;
        padding-bottom: 0.5rem;
        margin-top: 1.5rem;
    }
    .metric-card {
        background: #f8f9fa;
        border-left: 5px solid #0096DC;
        padding: 1rem;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .stButton>button {
        background-color: #0096DC;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.6rem 1.5rem;
        font-weight: 600;
    }
    .stButton>button:hover {
        background-color: #003366;
        color: white;
    }
    .report-expander {
        border: 1px solid #e0e0e0;
        border-radius: 10px;
        margin-bottom: 0.8rem;
    }
    .score-badge {
        font-size: 1.8rem;
        font-weight: 700;
        padding: 0.3rem 1rem;
        border-radius: 50px;
    }
    .stDataFrame {
        border-radius: 10px;
        overflow: hidden;
    }
    footer {
        text-align: center;
        color: #6c757d;
        font-size: 0.85rem;
        margin-top: 2rem;
    }
</style>
""", unsafe_allow_html=True)

# ====================== SESSION STATE ======================
def init_session_state():
    defaults = {
        "api_key": "",
        "jd_text": "",
        "jd_filename": "",
        "position_title": "",
        "cv_files": [],
        "cv_texts": {},
        "results": [],
        "client_name": "Apexire Client",
        "recruiter_name": "Apexire Recruitment Team",
        "generated_reports": {},
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val

init_session_state()

# ====================== HELPER FUNCTIONS ======================
def set_cell_shading(cell, color_hex: str):
    """Set background color for a table cell (e.g. '0096DC')"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def extract_text_from_uploaded_file(uploaded_file) -> str:
    """Robust text extraction for PDF, DOCX, TXT"""
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    try:
        if name.endswith(".pdf"):
            text = ""
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if not text.strip():
                # Fallback to pypdf
                reader = PdfReader(uploaded_file)
                for page in reader.pages:
                    text += page.extract_text() or ""
            return text.strip()
        elif name.endswith(".docx"):
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif name.endswith((".txt", ".text")):
            return uploaded_file.getvalue().decode("utf-8", errors="ignore")
        else:
            return f"[Unsupported file type: {uploaded_file.name}. Please use PDF, DOCX or TXT]"
    except Exception as e:
        return f"[Error extracting text from {uploaded_file.name}: {str(e)}]"

def get_openai_client(api_key: str) -> OpenAI:
    """Create OpenAI-compatible client for xAI Grok"""
    return OpenAI(
        api_key=api_key,
        base_url="https://api.x.ai/v1"  # xAI Grok API (OpenAI compatible)
    )

def evaluate_candidate(
    client: OpenAI,
    jd_text: str,
    cv_text: str,
    position_title: str,
    model: str = DEFAULT_MODEL
) -> Dict[str, Any]:
    """
    Call Grok to perform deep evaluation and return structured JSON.
    This prompt is engineered for the high-quality, consistent reports Apexire is known for.
    """
    system_prompt = f"""You are a world-class Senior Recruitment Consultant and Talent Strategist at Apexire, the leading AI-powered recruitment platform specializing in logistics, freight forwarding, supply chain, sales, warehouse, and professional services roles across Southern Africa and internationally.

Your reputation is built on **unbiased, evidence-based, highly actionable evaluations** that hiring managers trust. You never hallucinate experience. You are strict but fair. You highlight both strengths and genuine gaps with specific evidence from the CV.

**TASK**
Evaluate the CANDIDATE CV against the JOB DESCRIPTION provided below.

**OUTPUT FORMAT (MANDATORY)**
Return ONLY one valid, minified JSON object with exactly these keys (no markdown, no extra text, no ```json):

{{
  "candidate_name": "Exact full name from CV or 'Name Not Clearly Specified'",
  "contact_info": "Primary email and/or mobile number if present in CV. If none visible: 'Not provided in CV'",
  "position_applied": "{position_title}",
  "overall_match_score": <integer 0-100>,
  "executive_summary": "Write 3-5 crisp, professional sentences. Start with overall fit verdict. Highlight 1-2 strongest matches and 1 key concern or standout positive. Be specific.",
  "min_requirements_match": <integer 0-100>,
  "desirable_experience_match": <integer 0-100>,
  "experience_relevance": <integer 0-100>,
  "key_strengths": [
    "Specific, evidence-based strength #1 (quote relevant achievement or responsibility)",
    "Specific strength #2 with context",
    "... up to 7 items"
  ],
  "areas_for_development": [
    "Constructive, specific gap or area to probe in interview #1",
    "... 2 to 5 items. If truly excellent candidate, use: ['No significant development areas identified. Strong alignment across all criteria.']"
  ],
  "final_recommendation": {{
    "score": <integer 0-100>,
    "recommendation": "PROCEED (Strong Candidate) | PROCEED (Good Candidate) | CONSIDER (Average Fit) | REJECT (Poor Fit)",
    "rationale": "Write 4-7 sentences. Explain the holistic recommendation with specific references to JD requirements vs CV evidence. Mention interview focus areas, potential onboarding support needed, and overall business impact if hired. End with clear next-step recommendation."
  }}
}}

**SCORING RULES (be rigorous)**
- overall_match_score is a balanced holistic score (not simple average).
- min_requirements_match: How completely the candidate meets the MUST-HAVE / essential criteria listed in the JD.
- desirable_experience_match: Nice-to-have skills, advanced qualifications, extra tools/tech.
- experience_relevance: How directly transferable the candidate's past roles, industries, and responsibilities are to this specific position (even if title differs).
- High 90s = exceptional, clear top-tier fit with strong evidence.
- 80-89 = strong, minor gaps only.
- 70-79 = good but noticeable gaps or average relevance.
- 60-69 = average / borderline.
- Below 60 = significant mismatches or insufficient evidence.

**QUALITY STANDARDS**
- Every point in key_strengths and areas_for_development must be directly supported by text in the CV.
- Use quantified achievements where present (e.g. "Managed fleet of 47 trucks across 3 borders").
- Be honest about red flags (short tenure, career gaps, lack of required certifications, job-hopping without progression).
- For logistics/cross-border/fleet roles: pay special attention to compliance, cost control, multi-country experience, team leadership, systems (SAP, TMS, etc.).
- Tone: Professional, confident, objective, never overly salesy or harsh.

Now perform the evaluation."""

    user_prompt = f"""**JOB DESCRIPTION:**
{jd_text[:12000]}

**CANDIDATE CV:**
{cv_text[:15000]}

Evaluate now and return only the JSON object."""

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,  # Low temp for consistency and accuracy
            max_tokens=2800,
            response_format={"type": "json_object"}  # Enforce JSON if supported
        )
        content = response.choices[0].message.content.strip()
        
        # Clean possible markdown
        if content.startswith("```"):
            content = content.split("```")[1].strip()
            if content.lower().startswith("json"):
                content = content[4:].strip()
        
        data = json.loads(content)
        
        # Basic validation & defaults
        required = ["candidate_name", "contact_info", "position_applied", "overall_match_score",
                    "executive_summary", "min_requirements_match", "desirable_experience_match",
                    "experience_relevance", "key_strengths", "areas_for_development", "final_recommendation"]
        for k in required:
            if k not in data:
                data[k] = "N/A" if k != "key_strengths" else []
        
        if not isinstance(data.get("key_strengths"), list):
            data["key_strengths"] = [str(data["key_strengths"])]
        if not isinstance(data.get("areas_for_development"), list):
            data["areas_for_development"] = [str(data["areas_for_development"])]
            
        return data
        
    except Exception as e:
        # Fallback error object
        return {
            "candidate_name": "Evaluation Error",
            "contact_info": "N/A",
            "position_applied": position_title,
            "overall_match_score": 0,
            "executive_summary": f"Failed to evaluate due to technical issue: {str(e)}. Please try again or check CV/JD quality.",
            "min_requirements_match": 0,
            "desirable_experience_match": 0,
            "experience_relevance": 0,
            "key_strengths": ["Unable to complete analysis"],
            "areas_for_development": ["Retry evaluation or manual review recommended"],
            "final_recommendation": {
                "score": 0,
                "recommendation": "REJECT (Poor Fit)",
                "rationale": "Technical error during AI evaluation. Recommend manual review or re-upload clearer documents."
            }
        }

def create_beautiful_report(
    data: Dict[str, Any],
    logo_path: str,
    client_name: str,
    recruiter_name: str,
    output_path: str
) -> str:
    """
    Generate a premium, professional DOCX evaluation report.
    Significantly enhanced version of the provided template.
    """
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    # ========== HEADER ==========
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Cm(4)
    header_table.columns[1].width = Cm(14)
    
    # Logo cell
    logo_cell = header_table.rows[0].cells[0]
    if os.path.exists(logo_path):
        logo_para = logo_cell.paragraphs[0]
        run = logo_para.add_run()
        run.add_picture(logo_path, width=Cm(3.2))
    
    # Title cell
    title_cell = header_table.rows[0].cells[1]
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title_para.add_run("APEXIRE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = APEXIRE_BLUE
    
    subtitle = title_para.add_run("\nApexEval — Candidate Evaluation Report")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = APEXIRE_DARK
    
    # Date line
    date_para = title_cell.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(datetime.now().strftime("%d %B %Y"))
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # spacer
    
    # ========== CANDIDATE INFO TABLE ==========
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Table Grid'
    
    # Row 0: Candidate Name
    info_table.rows[0].cells[0].merge(info_table.rows[0].cells[1])
    cell = info_table.rows[0].cells[0]
    cell.text = "CANDIDATE NAME"
    set_cell_shading(cell, "003366")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            run.bold = True
    name_cell = info_table.rows[0].cells[2]
    name_cell.merge(info_table.rows[0].cells[3])
    name_cell.text = data.get("candidate_name", "N/A")
    for para in name_cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)
            run.bold = True
    
    # Row 1: Contact
    info_table.rows[1].cells[0].merge(info_table.rows[1].cells[1])
    cell = info_table.rows[1].cells[0]
    cell.text = "CONTACT"
    set_cell_shading(cell, "003366")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            run.bold = True
    contact_cell = info_table.rows[1].cells[2]
    contact_cell.merge(info_table.rows[1].cells[3])
    contact_cell.text = data.get("contact_info", "Not provided")
    
    # Row 2: Position + Score
    info_table.rows[2].cells[0].merge(info_table.rows[2].cells[1])
    cell = info_table.rows[2].cells[0]
    cell.text = "POSITION APPLIED"
    set_cell_shading(cell, "003366")
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            run.bold = True
    pos_cell = info_table.rows[2].cells[2]
    pos_cell.text = data.get("position_applied", "N/A")
    
    score_cell = info_table.rows[2].cells[3]
    score = data.get("overall_match_score", 0)
    score_cell.text = f"OVERALL MATCH\n{score}%"
    set_cell_shading(score_cell, "0096DC" if score >= 75 else ("FFC107" if score >= 60 else "DC3545"))
    for para in score_cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.bold = True
    
    doc.add_paragraph()
    
    # ========== EXECUTIVE SUMMARY ==========
    h = doc.add_paragraph()
    run = h.add_run("EXECUTIVE SUMMARY")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = APEXIRE_BLUE
    
    exec_sum = data.get("executive_summary", "")
    p = doc.add_paragraph(exec_sum)
    p.paragraph_format.space_after = Pt(8)
    
    # Recommendation line
    rec_data = data.get("final_recommendation", {})
    rec_text = rec_data.get("recommendation", "N/A")
    rec_para = doc.add_paragraph()
    run = rec_para.add_run("Recommendation: ")
    run.bold = True
    run.font.size = Pt(11)
    run2 = rec_para.add_run(rec_text)
    run2.bold = True
    run2.font.size = Pt(11)
    if "PROCEED (Strong" in rec_text:
        run2.font.color.rgb = SUCCESS_GREEN
    elif "PROCEED (Good" in rec_text:
        run2.font.color.rgb = RGBColor(0, 128, 0)
    elif "CONSIDER" in rec_text:
        run2.font.color.rgb = WARNING_YELLOW
    else:
        run2.font.color.rgb = DANGER_RED
    
    doc.add_paragraph()
    
    # ========== DETAILED BREAKDOWN ==========
    h = doc.add_paragraph()
    run = h.add_run("DETAILED BREAKDOWN")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = APEXIRE_BLUE
    
    breakdown_table = doc.add_table(rows=1, cols=3)
    breakdown_table.style = 'Table Grid'
    
    headers = ["Minimum Requirements", "Desirable Experience", "Experience Relevance"]
    values = [
        data.get("min_requirements_match", 0),
        data.get("desirable_experience_match", 0),
        data.get("experience_relevance", 0)
    ]
    
    for i, (header, value) in enumerate(zip(headers, values)):
        cell = breakdown_table.rows[0].cells[i]
        cell.text = f"{header}\n{value}%"
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(9)
            run.bold = True
        # Color coding
        if value >= 80:
            set_cell_shading(cell, "28A745")
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        elif value >= 65:
            set_cell_shading(cell, "FFC107")
        else:
            set_cell_shading(cell, "DC3545")
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    
    # ========== KEY STRENGTHS ==========
    h = doc.add_paragraph()
    run = h.add_run("KEY STRENGTHS")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = APEXIRE_BLUE
    
    strengths = data.get("key_strengths", [])
    if strengths:
        for s in strengths:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(s).font.size = Pt(10)
    else:
        doc.add_paragraph("No specific strengths highlighted in analysis.").italic = True
    
    # ========== AREAS FOR DEVELOPMENT ==========
    h = doc.add_paragraph()
    run = h.add_run("AREAS FOR DEVELOPMENT")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = APEXIRE_BLUE
    
    gaps = data.get("areas_for_development", [])
    if gaps and gaps != ["No significant development areas identified. Strong alignment across all criteria."]:
        for g in gaps:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(g).font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        run = p.add_run("✓ No significant development areas identified. Strong overall alignment.")
        run.font.color.rgb = SUCCESS_GREEN
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ========== FINAL RECOMMENDATION ==========
    h = doc.add_paragraph()
    run = h.add_run("FINAL RECOMMENDATION")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = APEXIRE_BLUE
    
    rationale = rec_data.get("rationale", "")
    p = doc.add_paragraph(rationale)
    p.paragraph_format.space_after = Pt(12)
    
    # ========== FOOTER ==========
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("━" * 60)
    run.font.color.rgb = RGBColor(200, 200, 200)
    
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        f"This evaluation has been prepared by {recruiter_name} for {client_name} using Apexire's AI-powered evaluation system (ApexEval).\n"
        "For internal recruitment use only. All assessments are evidence-based and designed to support fair, high-quality hiring decisions."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.italic = True
    
    # Save
    doc.save(output_path)
    return output_path

# ====================== UI ======================
def main():
    # Header
    col1, col2 = st.columns([1, 5])
    with col1:
        logo_path = "apexire_logo.jpg"
        if os.path.exists(logo_path):
            st.image(logo_path, width=85)
        else:
            st.markdown("**APEXIRE**")
    with col2:
        st.markdown(f"<h1 style='margin-bottom:0; color:#003366;'>{APP_NAME}</h1>", unsafe_allow_html=True)
        st.markdown(f"<p style='margin-top:0; color:#0096DC; font-size:1.1rem;'>{APP_TAGLINE}</p>", unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        api_key = st.text_input(
            "xAI API Key (Grok)",
            type="password",
            value=st.session_state.api_key,
            help="Get your key from https://console.x.ai/"
        )
        st.session_state.api_key = api_key
        
        model = st.selectbox(
            "Model",
            ["grok-3", "grok-3-mini", "grok-2"],
            index=0,
            help="grok-3 recommended for best reasoning quality"
        )
        
        st.divider()
        
        st.session_state.client_name = st.text_input("Client / Hiring Company", value=st.session_state.client_name)
        st.session_state.recruiter_name = st.text_input("Recruiter / Team Name", value=st.session_state.recruiter_name)
        
        st.divider()
        st.caption("📌 Upload your own TEMPLATE.docx in future versions for exact branding match.")
        st.caption("Current reports use an enhanced professional layout based on your template structure.")
        
        st.divider()
        st.markdown("**How to use**")
        st.markdown("""
        1. Paste your **xAI API key**
        2. Upload **Job Description** (PDF/DOCX)
        3. Upload up to **10 CVs**
        4. Click **Generate Reports**
        5. Review & download beautiful DOCX reports
        """)
    
    # Main content
    if not api_key:
        st.warning("🔑 Please enter your xAI API key in the sidebar to enable AI evaluation.")
        st.info("Don't have a key yet? Visit https://console.x.ai/ to create one (free tier available for testing).")
        return
    
    # Step 1: Job Description
    st.markdown("## 1. Job Description")
    jd_col1, jd_col2 = st.columns([2, 1])
    
    with jd_col1:
        jd_file = st.file_uploader(
            "Upload Job Description (PDF, DOCX or TXT)",
            type=["pdf", "docx", "txt"],
            key="jd_uploader",
            help="The JD will be analyzed in detail by Grok"
        )
        
        if jd_file:
            st.session_state.jd_filename = jd_file.name
            st.session_state.jd_text = extract_text_from_uploaded_file(jd_file)
            if not st.session_state.position_title:
                # Try to guess position from filename
                clean_name = jd_file.name.replace("_", " ").replace("-", " ").rsplit(".", 1)[0]
                st.session_state.position_title = clean_name.title()[:80]
    
    with jd_col2:
        st.session_state.position_title = st.text_input(
            "Position Title (auto-filled from filename)",
            value=st.session_state.position_title,
            placeholder="e.g. Cross Border Fleet Manager"
        )
    
    if st.session_state.jd_text:
        with st.expander("📄 Preview extracted JD text (first 800 chars)", expanded=False):
            st.text_area("", st.session_state.jd_text[:800] + "...", height=120, disabled=True)
    
    st.divider()
    
    # Step 2: CVs
    st.markdown("## 2. Candidate CVs (up to 10)")
    cv_files = st.file_uploader(
        "Drop CVs here (PDF or DOCX recommended)",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="cv_uploader",
        help=f"Maximum {MAX_CVS} files at a time for optimal processing"
    )
    
    if cv_files:
        if len(cv_files) > MAX_CVS:
            st.error(f"Maximum {MAX_CVS} CVs allowed. You uploaded {len(cv_files)}. Please reduce.")
            cv_files = cv_files[:MAX_CVS]
        
        st.session_state.cv_files = cv_files
        st.success(f"✅ {len(cv_files)} CV(s) ready for evaluation")
        
        # Show file list
        with st.expander("View uploaded files"):
            for f in cv_files:
                st.write(f"• {f.name} ({round(f.size/1024,1)} KB)")
    
    st.divider()
    
    # Generate Button
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        generate_btn = st.button(
            "🚀 Generate ApexEval Reports",
            type="primary",
            use_container_width=True,
            disabled=not (st.session_state.jd_text and st.session_state.cv_files and st.session_state.position_title)
        )
    
    if generate_btn:
        if not st.session_state.jd_text:
            st.error("Please upload a Job Description first.")
            return
        if not st.session_state.cv_files:
            st.error("Please upload at least one CV.")
            return
        if not st.session_state.position_title.strip():
            st.error("Please enter the Position Title.")
            return
        
        # Process
        client = get_openai_client(st.session_state.api_key)
        results = []
        progress_bar = st.progress(0, text="Initializing evaluation...")
        status_container = st.status("Evaluating candidates with Grok...", expanded=True)
        
        total = len(st.session_state.cv_files)
        
        for idx, cv_file in enumerate(st.session_state.cv_files):
            progress = (idx + 1) / total
            progress_bar.progress(progress, text=f"Processing {idx+1}/{total}: {cv_file.name}")
            
            with status_container:
                st.write(f"🔍 Analyzing **{cv_file.name}** ...")
            
            cv_text = extract_text_from_uploaded_file(cv_file)
            if not cv_text or len(cv_text) < 50:
                st.warning(f"Very little text extracted from {cv_file.name}. Results may be limited.")
            
            eval_data = evaluate_candidate(
                client=client,
                jd_text=st.session_state.jd_text,
                cv_text=cv_text,
                position_title=st.session_state.position_title,
                model=model
            )
            
            # Attach original filename for reference
            eval_data["_source_file"] = cv_file.name
            results.append(eval_data)
        
        progress_bar.progress(1.0, text="All evaluations complete!")
        status_container.update(label="✅ All reports generated successfully", state="complete")
        
        st.session_state.results = results
        st.session_state.generated_reports = {}  # reset
        
        st.balloons()
        st.success(f"🎉 Successfully evaluated {len(results)} candidate(s)! Scroll down to review and download reports.")
    
    # Results Section
    if st.session_state.results:
        st.markdown("---")
        st.markdown("## 📊 Evaluation Results Dashboard")
        
        results = st.session_state.results
        
        # Summary metrics
        scores = [r.get("overall_match_score", 0) for r in results]
        avg_score = sum(scores) / len(scores) if scores else 0
        top_score = max(scores) if scores else 0
        strong_candidates = sum(1 for r in results if "PROCEED (Strong" in str(r.get("final_recommendation", {}).get("recommendation", "")))
        
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Candidates Evaluated", len(results))
        m2.metric("Average Match Score", f"{avg_score:.0f}%")
        m3.metric("Highest Score", f"{top_score}%")
        m4.metric("Strong Recommendations", f"{strong_candidates} / {len(results)}")
        
        # Summary Table
        st.markdown("### Quick Overview")
        summary_data = []
        for r in results:
            rec = r.get("final_recommendation", {}).get("recommendation", "N/A")
            summary_data.append({
                "Candidate": r.get("candidate_name", "Unknown")[:35],
                "Source File": r.get("_source_file", "")[:30],
                "Overall Score": r.get("overall_match_score", 0),
                "Min Req": r.get("min_requirements_match", 0),
                "Desirable": r.get("desirable_experience_match", 0),
                "Relevance": r.get("experience_relevance", 0),
                "Recommendation": rec.replace(" (Strong Candidate)", "").replace(" (Good Candidate)", "").replace(" (Average Fit)", "").replace(" (Poor Fit)", "")
            })
        
        import pandas as pd
        df = pd.DataFrame(summary_data)
        st.dataframe(
            df.style.background_gradient(subset=["Overall Score", "Min Req", "Desirable", "Relevance"], cmap="RdYlGn", vmin=40, vmax=100),
            use_container_width=True,
            hide_index=True
        )
        
        st.divider()
        
        # Individual Reports
        st.markdown("### Detailed Reports")
        st.caption("Expand each candidate to view the full AI-generated evaluation. Download individual DOCX reports below each preview.")
        
        logo_path = "apexire_logo.jpg"
        
        for i, data in enumerate(results):
            name = data.get("candidate_name", f"Candidate {i+1}")
            score = data.get("overall_match_score", 0)
            rec = data.get("final_recommendation", {}).get("recommendation", "")
            
            # Color indicator
            if score >= 80:
                indicator = "🟢"
            elif score >= 65:
                indicator = "🟡"
            else:
                indicator = "🔴"
            
            with st.expander(f"{indicator} **{name}** — {score}% | {rec}", expanded=(i == 0)):
                # Quick view in markdown
                st.markdown(f"**Executive Summary**")
                st.write(data.get("executive_summary", ""))
                
                st.markdown(f"**Key Strengths** ({len(data.get('key_strengths', []))})")
                for s in data.get("key_strengths", []):
                    st.markdown(f"- {s}")
                
                st.markdown(f"**Areas for Development**")
                for g in data.get("areas_for_development", []):
                    st.markdown(f"- {g}")
                
                st.markdown(f"**Final Rationale**")
                st.write(data.get("final_recommendation", {}).get("rationale", ""))
                
                # Download button for this report
                if st.button(f"📥 Download {name[:25]} Report (DOCX)", key=f"dl_{i}"):
                    with st.spinner("Generating professional DOCX report..."):
                        safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()[:40]
                        filename = f"ApexEval_{safe_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                        
                        tmp_path = os.path.join(tempfile.gettempdir(), filename)
                        create_beautiful_report(
                            data=data,
                            logo_path=logo_path,
                            client_name=st.session_state.client_name,
                            recruiter_name=st.session_state.recruiter_name,
                            output_path=tmp_path
                        )
                        
                        with open(tmp_path, "rb") as f:
                            st.download_button(
                                label="✅ Click to Download Report",
                                data=f,
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_btn_{i}"
                            )
        
        # Bulk Download
        st.divider()
        st.markdown("### 📦 Bulk Actions")
        
        if st.button("📥 Download ALL Reports as ZIP", type="secondary", use_container_width=True):
            with st.spinner("Creating ZIP archive of all professional reports..."):
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for i, data in enumerate(results):
                        safe_name = "".join(c for c in data.get("candidate_name", f"Candidate_{i}") if c.isalnum() or c in (' ', '-', '_')).rstrip()[:35]
                        filename = f"ApexEval_{safe_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                        
                        tmp_path = os.path.join(tempfile.gettempdir(), f"temp_report_{i}.docx")
                        create_beautiful_report(
                            data=data,
                            logo_path=logo_path,
                            client_name=st.session_state.client_name,
                            recruiter_name=st.session_state.recruiter_name,
                            output_path=tmp_path
                        )
                        zipf.write(tmp_path, arcname=filename)
                
                zip_buffer.seek(0)
                st.download_button(
                    label="✅ Download Complete ZIP Archive",
                    data=zip_buffer,
                    file_name=f"ApexEval_Reports_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
                    mime="application/zip"
                )
        
        st.caption("All reports are generated using an enhanced professional template aligned with your original structure. Future versions will support uploading your exact TEMPLATE.docx for pixel-perfect branding.")

# Run
if __name__ == "__main__":
    main()
